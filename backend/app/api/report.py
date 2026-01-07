from fastapi import APIRouter, HTTPException, Depends
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import json
import os
import io
import math
import numpy as np
import hashlib

# Monkey patch for hashlib.md5 if usedforsecurity is not supported (ReportLab compatibility)
try:
    hashlib.md5(b'', usedforsecurity=False)
except TypeError:
    _original_md5 = hashlib.md5
    def _md5_wrapper(string=b'', usedforsecurity=True):
        return _original_md5(string)
    hashlib.md5 = _md5_wrapper

import matplotlib
# Use Agg backend for non-interactive plotting
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from fastapi.responses import StreamingResponse
from app.api import deps
import pandas as pd

router = APIRouter()

class ReportRequest(BaseModel):
    workflow_name: str
    format: str = "docx"  # "docx" or "pdf"

def get_workflow_path(name: str, user_dir: str) -> str:
    return os.path.join(user_dir, "workflows", f"{name}.json")

def get_meta_path(filename: str, user_dir: str) -> str:
    return os.path.join(user_dir, "uploads", f"{filename}.meta.json")

def configure_matplotlib_chinese():
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
    plt.rcParams['axes.unicode_minus'] = False

def extract_workflow_data(workflow_data: Dict, user_dir: str):
    nodes = workflow_data.get("nodes", [])
    edges = workflow_data.get("edges", [])
    
    # 1. Extract EDA Data
    eda_stats = None
    data_filename = None
    data_node = next((n for n in nodes if n.get("type") == "dataNode"), None)
    if data_node and data_node.get("data", {}).get("filename"):
        data_filename = data_node["data"]["filename"]
        meta_path = get_meta_path(data_filename, user_dir)
        if os.path.exists(meta_path):
            with open(meta_path, 'r', encoding='utf-8') as f:
                eda_stats = json.load(f)

    # 2. Extract Algorithm Results
    algo_results = []
    
    # Helper to find connected EvalNode
    def find_eval_node_result(algo_node_id):
        # BFS to find downstream evalNode
        queue = [algo_node_id]
        visited = set()
        while queue:
            curr = queue.pop(0)
            if curr in visited: continue
            visited.add(curr)
            
            # Find outgoing edges
            outgoing = [e for e in edges if e["source"] == curr]
            for e in outgoing:
                target_id = e["target"]
                target_node = next((n for n in nodes if n["id"] == target_id), None)
                if target_node:
                    if target_node.get("type") == "evalNode" and target_node.get("data", {}).get("result"):
                        return target_node["data"]["result"]
                    queue.append(target_id)
        return None

    algo_nodes = [n for n in nodes if n.get("type") == "algoNode" and n.get("data", {}).get("category") == "Model"]
    
    for node in algo_nodes:
        result = None
        # Check internal result first
        if node.get("data", {}).get("result"):
            result = node["data"]["result"]
        else:
            # Check downstream eval node
            result = find_eval_node_result(node["id"])
            
        if result:
            algo_results.append({
                "id": node["id"],
                "label": node["data"].get("label", "Unknown Algo"),
                "category": node["data"].get("category"),
                "type": result.get("type"), # classification, regression, clustering
                "result": result,
                "params": node["data"].get("params", {})
            })
            
    return data_filename, eda_stats, algo_results

def generate_charts(algo_results: List[Dict]):
    configure_matplotlib_chinese()
    charts = {}
    
    # Classification: Radar Chart
    clf_results = [r for r in algo_results if r["type"] == "classification"]
    if clf_results:
        # Prepare data
        labels = ['Accuracy', 'Precision', 'Recall', 'F1']
        
        fig = plt.figure(figsize=(8, 6))
        ax = fig.add_subplot(111, polar=True)
        
        # Determine range
        values_list = []
        for res in clf_results:
            report = res["result"].get("report", {}).get("weighted avg", {})
            vals = [
                res["result"].get("accuracy", 0),
                report.get("precision", 0),
                report.get("recall", 0),
                report.get("f1-score", 0)
            ]
            values_list.append(vals)
            
            # Close the loop
            vals_plot = vals + [vals[0]]
            angles = np.linspace(0, 2*np.pi, len(labels), endpoint=False).tolist()
            angles_plot = angles + [angles[0]]
            
            ax.plot(angles_plot, vals_plot, linewidth=2, label=res["label"])
            ax.fill(angles_plot, vals_plot, alpha=0.25)
            
        ax.set_thetagrids(np.degrees(angles), labels)
        plt.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1))
        plt.title("分类模型性能对比 (Radar Chart)")
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        charts['classification_radar'] = buf
        plt.close(fig)

    # Regression: Bar Chart (MAE)
    reg_results = [r for r in algo_results if r["type"] == "regression"]
    if reg_results:
        labels = [r["label"] for r in reg_results]
        mae_values = [r["result"].get("mae", 0) for r in reg_results]
        
        fig, ax = plt.subplots(figsize=(8, 5))
        bars = ax.bar(labels, mae_values, color='skyblue')
        
        ax.set_ylabel('MAE (Mean Absolute Error)')
        ax.set_title('回归模型误差对比 (越小越好)')
        
        # Add value labels
        for bar in bars:
            height = bar.get_height()
            ax.annotate(f'{height:.4f}',
                        xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 3),  # 3 points vertical offset
                        textcoords="offset points",
                        ha='center', va='bottom')
                        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        charts['regression_bar'] = buf
        plt.close(fig)

    # Clustering: Bar Chart (Silhouette Score)
    clus_results = [r for r in algo_results if r["type"] == "clustering"]
    if clus_results:
        labels = [r["label"] for r in clus_results]
        sil_values = [r["result"].get("silhouette_score", -1) for r in clus_results]
        
        fig, ax = plt.subplots(figsize=(8, 5))
        bars = ax.bar(labels, sil_values, color='lightgreen')
        
        ax.set_ylabel('Silhouette Score (轮廓系数)')
        ax.set_title('聚类模型轮廓系数对比 (越大越好)')
        
        # Add value labels
        for bar in bars:
            height = bar.get_height()
            ax.annotate(f'{height:.4f}',
                        xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 3),
                        textcoords="offset points",
                        ha='center', va='bottom')
                        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        charts['clustering_bar'] = buf
        plt.close(fig)
        
    return charts

def create_docx(workflow_name: str, data_filename: str, eda_stats: Dict, algo_results: List[Dict], charts: Dict):
    doc = Document()
    
    # Set style for Chinese font support (Basic setup)
    try:
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    except Exception as e:
        print(f"Warning: Failed to set docx font: {e}")
    
    # Title
    heading = doc.add_heading(f'算法分析报告 - {workflow_name}', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'生成时间: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'数据来源: {data_filename or "未指定"}')
    
    # 1. EDA Section
    doc.add_heading('1. 输入数据探索性分析 (EDA)', level=1)
    
    if eda_stats:
        doc.add_paragraph(f"总行数: {eda_stats.get('rows', 0)}")
        cols = eda_stats.get("columns", {})
        
        # Create Table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_labels = ['列名', '类型', '缺失值', '缺失比例(%)', '异常值', '唯一值']
        for i, label in enumerate(hdr_labels):
            hdr_cells[i].text = label
            
        for col_name, stats in cols.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(col_name)
            row_cells[1].text = str(stats.get('type', ''))
            row_cells[2].text = str(stats.get('missing', 0))
            row_cells[3].text = str(stats.get('missing_pct', 0))
            row_cells[4].text = str(stats.get('outliers', 0))
            row_cells[5].text = str(stats.get('unique', 0))
    else:
        doc.add_paragraph("暂无 EDA 数据。请确保在工作流中运行了数据加载节点。")
        
    # 2. Algorithm Results Section
    doc.add_heading('2. 算法可视化结果分析', level=1)
    
    if not algo_results:
        doc.add_paragraph("暂无算法运行结果。")
    else:
        # 2.1 Visualization
        doc.add_heading('2.1 可视化对比', level=2)
        
        if 'classification_radar' in charts:
            doc.add_paragraph('分类模型性能雷达图：')
            doc.add_picture(charts['classification_radar'], width=Inches(5.5))
            
        if 'regression_bar' in charts:
            doc.add_paragraph('回归模型误差对比图：')
            doc.add_picture(charts['regression_bar'], width=Inches(5.5))

        if 'clustering_bar' in charts:
            doc.add_paragraph('聚类模型轮廓系数对比图：')
            doc.add_picture(charts['clustering_bar'], width=Inches(5.5))
            
        # 2.2 Metrics Table
        doc.add_heading('2.2 详细指标数据', level=2)
        
        # Classification Table
        clf_results = [r for r in algo_results if r["type"] == "classification"]
        if clf_results:
            doc.add_heading('分类模型指标', level=3)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = '算法'
            hdr[1].text = 'Accuracy'
            hdr[2].text = 'Precision (Weighted)'
            hdr[3].text = 'Recall (Weighted)'
            hdr[4].text = 'F1 (Weighted)'
            
            for res in clf_results:
                row = table.add_row().cells
                report = res["result"].get("report", {}) or {}
                weighted_avg = report.get("weighted avg", {}) or {}
                
                row[0].text = res["label"]
                row[1].text = f"{res['result'].get('accuracy', 0):.4f}"
                row[2].text = f"{weighted_avg.get('precision', 0):.4f}"
                row[3].text = f"{weighted_avg.get('recall', 0):.4f}"
                row[4].text = f"{weighted_avg.get('f1-score', 0):.4f}"
                
        # Regression Table
        reg_results = [r for r in algo_results if r["type"] == "regression"]
        if reg_results:
            doc.add_heading('回归模型指标', level=3)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = '算法'
            hdr[1].text = 'MAE'
            hdr[2].text = 'MSE'
            hdr[3].text = 'R2 Score'
            
            for res in reg_results:
                row = table.add_row().cells
                row[0].text = res["label"]
                row[1].text = f"{res['result'].get('mae', 0):.4f}"
                row[2].text = f"{res['result'].get('mse', 0):.4f}"
                row[3].text = f"{res['result'].get('r2_score', 0):.4f}"

        # Clustering Table
        clus_results = [r for r in algo_results if r["type"] == "clustering"]
        if clus_results:
            doc.add_heading('聚类模型指标', level=3)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = '算法'
            hdr[1].text = 'Clusters'
            hdr[2].text = 'Silhouette'
            hdr[3].text = 'Calinski-Harabasz'
            hdr[4].text = 'Davies-Bouldin'
            
            for res in clus_results:
                row = table.add_row().cells
                row[0].text = res["label"]
                row[1].text = str(res['result'].get('n_clusters', 0))
                row[2].text = f"{res['result'].get('silhouette_score', -1):.4f}"
                row[3].text = f"{res['result'].get('calinski_harabasz_score', -1):.4f}"
                row[4].text = f"{res['result'].get('davies_bouldin_score', -1):.4f}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_pdf(workflow_name: str, data_filename: str, eda_stats: Dict, algo_results: List[Dict], charts: Dict):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    elements = []
    
    # Register Font
    try:
        pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc')) # Windows default
        font_name = 'SimSun'
    except:
        try:
             pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
             font_name = 'SimHei'
        except:
             # Fallback to standard font (Chinese will be broken)
             font_name = 'Helvetica'
    
    styles = getSampleStyleSheet()
    # Define Chinese Style
    cn_title_style = ParagraphStyle(name='CNTitle', parent=styles['Title'], fontName=font_name, fontSize=18, leading=22)
    cn_h1_style = ParagraphStyle(name='CNHeading1', parent=styles['Heading1'], fontName=font_name, fontSize=16, leading=20, spaceBefore=12, spaceAfter=6)
    cn_h2_style = ParagraphStyle(name='CNHeading2', parent=styles['Heading2'], fontName=font_name, fontSize=14, leading=18, spaceBefore=10, spaceAfter=6)
    cn_body_style = ParagraphStyle(name='CNBody', parent=styles['Normal'], fontName=font_name, fontSize=10, leading=14)
    
    # Title
    elements.append(Paragraph(f'算法分析报告 - {workflow_name}', cn_title_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f'生成时间: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}', cn_body_style))
    elements.append(Paragraph(f'数据来源: {data_filename or "未指定"}', cn_body_style))
    elements.append(Spacer(1, 24))
    
    # 1. EDA
    elements.append(Paragraph('1. 输入数据探索性分析 (EDA)', cn_h1_style))
    if eda_stats:
        elements.append(Paragraph(f"总行数: {eda_stats.get('rows', 0)}", cn_body_style))
        elements.append(Spacer(1, 12))
        
        cols = eda_stats.get("columns", {})
        data = [['列名', '类型', '缺失值', '缺失比(%)', '异常值', '唯一值']]
        for col_name, stats in cols.items():
            data.append([
                col_name[:15], # Truncate long names
                stats.get('type', ''),
                str(stats.get('missing', 0)),
                str(stats.get('missing_pct', 0)),
                str(stats.get('outliers', 0)),
                str(stats.get('unique', 0))
            ])
            
        t = Table(data)
        t.setStyle(TableStyle([
            ('FONTNAME', (0,0), (-1,-1), font_name),
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(t)
    else:
        elements.append(Paragraph("暂无 EDA 数据。", cn_body_style))
        
    elements.append(Spacer(1, 24))
    
    # 2. Results
    elements.append(Paragraph('2. 算法可视化结果分析', cn_h1_style))
    
    if not algo_results:
        elements.append(Paragraph("暂无算法运行结果。", cn_body_style))
    else:
        elements.append(Paragraph('2.1 可视化对比', cn_h2_style))
        
        if 'classification_radar' in charts:
            elements.append(Paragraph('分类模型性能雷达图：', cn_body_style))
            elements.append(RLImage(charts['classification_radar'], width=400, height=300))
            elements.append(Spacer(1, 12))
            
        if 'regression_bar' in charts:
            elements.append(Paragraph('回归模型误差对比图：', cn_body_style))
            elements.append(RLImage(charts['regression_bar'], width=400, height=250))
            elements.append(Spacer(1, 12))

        if 'clustering_bar' in charts:
            elements.append(Paragraph('聚类模型轮廓系数对比图：', cn_body_style))
            elements.append(RLImage(charts['clustering_bar'], width=400, height=250))
            elements.append(Spacer(1, 12))
            
        elements.append(PageBreak())
        elements.append(Paragraph('2.2 详细指标数据', cn_h2_style))
        
        # Clf Table
        clf_results = [r for r in algo_results if r["type"] == "classification"]
        if clf_results:
            elements.append(Paragraph('分类模型指标', cn_body_style))
            data = [['算法', 'Acc', 'Prec', 'Recall', 'F1']]
            for res in clf_results:
                report = res["result"].get("report", {}) or {}
                weighted_avg = report.get("weighted avg", {}) or {}
                data.append([
                    res["label"],
                    f"{res['result'].get('accuracy', 0):.4f}",
                    f"{weighted_avg.get('precision', 0):.4f}",
                    f"{weighted_avg.get('recall', 0):.4f}",
                    f"{weighted_avg.get('f1-score', 0):.4f}"
                ])
            t = Table(data)
            t.setStyle(TableStyle([
                ('FONTNAME', (0,0), (-1,-1), font_name),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 12))
            
        # Reg Table
        reg_results = [r for r in algo_results if r["type"] == "regression"]
        if reg_results:
            elements.append(Paragraph('回归模型指标', cn_body_style))
            data = [['算法', 'MAE', 'MSE', 'R2']]
            for res in reg_results:
                data.append([
                    res["label"],
                    f"{res['result'].get('mae', 0):.4f}",
                    f"{res['result'].get('mse', 0):.4f}",
                    f"{res['result'].get('r2_score', 0):.4f}"
                ])
            t = Table(data)
            t.setStyle(TableStyle([
                ('FONTNAME', (0,0), (-1,-1), font_name),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 12))

        # Clustering Table
        clus_results = [r for r in algo_results if r["type"] == "clustering"]
        if clus_results:
            elements.append(Paragraph('聚类模型指标', cn_body_style))
            data = [['算法', 'Clusters', 'Silhouette', 'Calinski', 'Davies']]
            for res in clus_results:
                data.append([
                    res["label"],
                    str(res['result'].get('n_clusters', 0)),
                    f"{res['result'].get('silhouette_score', -1):.4f}",
                    f"{res['result'].get('calinski_harabasz_score', -1):.4f}",
                    f"{res['result'].get('davies_bouldin_score', -1):.4f}"
                ])
            t = Table(data)
            t.setStyle(TableStyle([
                ('FONTNAME', (0,0), (-1,-1), font_name),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ]))
            elements.append(t)

            
    doc.build(elements)
    buf.seek(0)
    return buf

@router.post("/export")
async def export_report(
    request: ReportRequest,
    user_dir: str = Depends(deps.get_current_user_dir)
):
    try:
        # 1. Load Workflow
        workflow_path = get_workflow_path(request.workflow_name, user_dir)
        if not os.path.exists(workflow_path):
            raise HTTPException(status_code=404, detail="Workflow not found")
            
        with open(workflow_path, 'r', encoding='utf-8') as f:
            workflow_data = json.load(f)
            
        # 2. Extract Data
        data_filename, eda_stats, algo_results = extract_workflow_data(workflow_data, user_dir)
        
        # 3. Generate Charts
        charts = generate_charts(algo_results)
        
        # 4. Generate File
        if request.format == "docx":
            file_stream = create_docx(request.workflow_name, data_filename, eda_stats, algo_results, charts)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{request.workflow_name}_report.docx"
        elif request.format == "pdf":
            file_stream = create_pdf(request.workflow_name, data_filename, eda_stats, algo_results, charts)
            media_type = "application/pdf"
            filename = f"{request.workflow_name}_report.pdf"
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")
            
        return StreamingResponse(
            file_stream, 
            media_type=media_type, 
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
